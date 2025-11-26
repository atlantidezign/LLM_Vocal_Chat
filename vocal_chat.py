import gradio as gr
import speech_recognition as sr
import pyttsx3
import requests
import json
import threading
from datetime import datetime
import os
import numpy as np
from scipy import signal

# Whisper import (handles both installed and not installed cases)
try:
    import whisper
    import torch
    WHISPER_AVAILABLE = True
    CUDA_AVAILABLE = torch.cuda.is_available()
    GPU_NAME = torch.cuda.get_device_name(0) if CUDA_AVAILABLE else "N/A"
    whisper_model = None
except ImportError:
    WHISPER_AVAILABLE = False
    CUDA_AVAILABLE = False
    GPU_NAME = "N/A"
    whisper_model = None

# Librosa import for audio resampling (optional)
try:
    import librosa
    LIBROSA_AVAILABLE = True
except ImportError:
    LIBROSA_AVAILABLE = False

# Python-docx import for DOCX export
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ===== LOCALIZATION =====
TRANSLATIONS = {
    "en": {
        # Main UI
        "title": "🎙️ Voice Chat with LM Studio",
        "subtitle": "Interact with AI models via voice or text",
        "conversation": "Conversation",
        "record_audio": "🎤 Record Audio",
        "text_input": "✍️ Text (auto-filled from audio or type manually)",
        "text_placeholder": "Record audio or type here...",
        "send": "📤 Send",
        "stop_audio": "⏹️ Stop Audio",
        "clear_chat": "🗑️ Clear Chat",
        
        # Quick settings
        "quick_settings": "⚙️ Quick Settings",
        "connection_status": "Verifying connection...",
        "voice_recognition": "🎤 Voice Recognition",
        "recognition_method": "Recognition Method",
        "google_option": "google - Online (fast, requires internet)",
        "whisper_option": "whisper - Local (more accurate, privacy)",
        "audio_language": "🌍 Audio Input Language",
        "whisper_available": "✅ **Whisper available**",
        "gpu_acceleration": "🎮 GPU Acceleration Active",
        "whisper_model": "Whisper Model",
        "whisper_sizes": "tiny=39MB, base=74MB, small=244MB, medium=769MB, large=1.5GB (recommended with GPU!)",
        "use_gpu": "🎮 Use GPU",
        "load_model": "📥 Load Model",
        "model_status": "Status",
        "whisper_not_installed": "⚠️ **Whisper not installed**\n\nTo use local recognition:\n```bash\npip install openai-whisper\n```",
        "speak_response": "🔊 Read response aloud",
        "ui_language": "🌐 UI Language",
        
        # Model parameters
        "model_params": "🎛️ Model Parameters",
        "model_label": "🤖 Model",
        "model_info": "Select model from LM Studio",
        "temperature_label": "🌡️ Temperature (Creativity)",
        "max_tokens_label": "📏 Max Tokens (Length)",
        
        # Response customization
        "response_custom": "🎨 Response Customization",
        "system_prompt": "📝 System Prompt",
        "system_prompt_default": "You are a helpful assistant that responds concisely and clearly.",
        
        # Voice settings
        "voice_settings": "🔊 Voice Settings",
        "voice_label": "🎭 Voice",
        "speed_label": "⚡ Speed",
        "volume_label": "🔊 Volume",
        
        # Export
        "export_conv": "💾 Export Conversation",
        "save_json": "💾 Save JSON",
        "export_txt": "📄 Export TXT",
        "generate_summary": "📝 Generate Summary (MD + DOCX)",
        "save_status": "Save Status",
        "summary_preview": "📄 Summary Preview",
        "files_saved": "ℹ️ **Files saved** - Check project folder for .md and .docx files",
        
        # Advanced settings
        "advanced_settings": "🔧 Advanced Settings",
        "lm_config": "🌐 LM Studio Configuration",
        "lm_url": "LM Studio Server URL",
        "test_connection": "🔌 Test Connection",
        "connection_result": "Connection Status",
        "gpu_info": "🎮 GPU Info",
        "gpu_detected": "✅ **GPU Detected**",
        "gpu_not_detected": "⚠️ **GPU not detected** - Using CPU",
        "gpu_recommendations": "**Recommendations for RTX 3090:**\n- Use Whisper `large` model for max accuracy\n- Enable GPU acceleration in LM Studio\n- Automatic FP16 active for speed",
        "system_info": "📋 System Info",
        
        # Guide
        "guide": "ℹ️ Guide",
        "guide_content": """## 🎯 How to use the application

### 🎤 Voice Recognition

**Google (Online) - Default**
- ✅ Fast and reliable
- ✅ No installation required
- ❌ Requires internet connection
- ❌ Audio data sent to Google

**Whisper (Local) - Recommended for Privacy**
- ✅ Completely local and private
- ✅ Much more accurate
- ✅ Supports 99+ languages
- ✅ Better with accents and noise
- ❌ Requires installation: `pip install openai-whisper`
- ❌ Slightly slower

**Whisper Models:**
- `tiny` (39 MB): Very fast, basic accuracy
- `base` (74 MB): Good for CPU
- `small` (244 MB): Good compromise
- `medium` (769 MB): **Recommended for GPU** - Great quality/speed ratio
- `large` (1.5 GB): **Maximum accuracy with RTX 3090** - Fast on powerful GPUs!

### 🎮 GPU Acceleration
With your RTX 3090:
- Whisper `large` is very fast (~2-3 seconds per audio)
- Automatic FP16 for optimal performance
- Recommended: use `medium` or `large`

### 📝 Summary Document
At the end of conversation, click **"📝 Generate Summary (MD + DOCX)"** to:
- Automatically create a structured markdown document
- Summarize topics, decisions and action items
- Save as downloadable .md and .docx files
- View preview before download

### 🎤 Voice Input
1. Click the **microphone** 🎤
2. **Speak** your message
3. Click **STOP** ⏹️ to end recording
4. ⏳ *App transcribes automatically...*
5. ✅ **Text appears in editable field**
6. 📝 **You can edit/correct** text if needed
7. Click **"📤 Send"** to send to LM Studio
8. 🤖 Receive response (written and/or voiced)

**Advantages:**
- ✅ See what was understood before sending
- ✅ Correct transcription errors
- ✅ Add details or modifications
- ✅ Combine audio + manual editing
""",
        
        # Status messages
        "no_text": "❌ No text to send. Record audio and wait for transcription, then send.",
        "session_restored": "✅ Last session loaded",
        "conv_reset": "🔄 Conversation reset",
        "response_time": "📊 Response time",
        "total_messages": "Total messages",
        "avg_time": "Average",
        "playing_audio": "🔊 Playing audio...",
        "response_received": "✅ Response received",
        "no_conversation": "❌ No conversation to",
        "json_saved": "✅ Conversation exported to",
        "txt_saved": "✅ Conversation exported to",
        "summary_error": "❌ Error generating summary",
        "connected": "✅ Connected",
        "not_connected": "❌ Cannot connect to LM Studio",
        "models_available": "models available",
        "model_loaded": "✅ Model loaded",
        "cached": "✅ Loaded from cache",
        "downloaded": "✅ Download completed and cached",
        "cache_location": "📁 Cache",
    },
    "it": {
        # Main UI
        "title": "🎙️ Chat Vocale con LM Studio",
        "subtitle": "Interagisci con i modelli AI tramite voce o testo",
        "conversation": "Conversazione",
        "record_audio": "🎤 Registra Audio",
        "text_input": "✍️ Testo (auto-compilato da audio o scrivi manualmente)",
        "text_placeholder": "Registra audio o scrivi qui...",
        "send": "📤 Invia",
        "stop_audio": "⏹️ Ferma Audio",
        "clear_chat": "🗑️ Pulisci Chat",
        
        # Quick settings  
        "quick_settings": "⚙️ Impostazioni Rapide",
        "connection_status": "Verifica connessione...",
        "voice_recognition": "🎤 Riconoscimento Vocale",
        "recognition_method": "Metodo di Riconoscimento",
        "google_option": "google - Online (veloce, richiede internet)",
        "whisper_option": "whisper - Locale (più accurato, privacy)",
        "audio_language": "🌍 Lingua Audio Input",
        "whisper_available": "✅ **Whisper disponibile**",
        "gpu_acceleration": "🎮 Accelerazione GPU Attiva",
        "whisper_model": "Modello Whisper",
        "whisper_sizes": "tiny=39MB, base=74MB, small=244MB, medium=769MB, large=1.5GB (consigliato con GPU!)",
        "use_gpu": "🎮 Usa GPU",
        "load_model": "📥 Carica Modello",
        "model_status": "Status",
        "whisper_not_installed": "⚠️ **Whisper non installato**\n\nPer usare il riconoscimento locale:\n```bash\npip install openai-whisper\n```",
        "speak_response": "🔊 Leggi risposta ad alta voce",
        "ui_language": "🌐 Lingua Interfaccia",
        
        # Model parameters
        "model_params": "🎛️ Parametri Modello",
        "model_label": "🤖 Modello",
        "model_info": "Seleziona il modello da LM Studio",
        "temperature_label": "🌡️ Temperature (Creatività)",
        "max_tokens_label": "📏 Max Tokens (Lunghezza)",
        
        # Response customization
        "response_custom": "🎨 Personalizzazione Risposta",
        "system_prompt": "📝 System Prompt",
        "system_prompt_default": "Sei un assistente utile che risponde in modo conciso e chiaro.",
        
        # Voice settings
        "voice_settings": "🔊 Impostazioni Voce",
        "voice_label": "🎭 Voce",
        "speed_label": "⚡ Velocità",
        "volume_label": "🔊 Volume",
        
        # Export
        "export_conv": "💾 Esporta Conversazione",
        "save_json": "💾 Salva JSON",
        "export_txt": "📄 Esporta TXT",
        "generate_summary": "📝 Genera Summary (MD + DOCX)",
        "save_status": "Status Salvataggio",
        "summary_preview": "📄 Anteprima Summary",
        "files_saved": "ℹ️ **File salvati** - Controlla la cartella del progetto per i file .md e .docx",
        
        # Advanced settings
        "advanced_settings": "🔧 Impostazioni Avanzate",
        "lm_config": "🌐 Configurazione LM Studio",
        "lm_url": "URL Server LM Studio",
        "test_connection": "🔌 Test Connessione",
        "connection_result": "Status Connessione",
        "gpu_info": "🎮 Info GPU",
        "gpu_detected": "✅ **GPU Rilevata**",
        "gpu_not_detected": "⚠️ **GPU non rilevata** - Utilizzo CPU",
        "gpu_recommendations": "**Raccomandazioni per RTX 3090:**\n- Usa modello Whisper `large` per massima accuratezza\n- Abilita accelerazione GPU in LM Studio\n- FP16 automatico attivo per velocità",
        "system_info": "📋 Info Sistema",
        
        # Guide
        "guide": "ℹ️ Guida",
        "guide_content": """## 🎯 Come usare l'applicazione

### 🎤 Riconoscimento Vocale

**Google (Online) - Default**
- ✅ Veloce e affidabile
- ✅ Non richiede installazione
- ❌ Richiede connessione internet
- ❌ I dati audio vengono inviati a Google

**Whisper (Locale) - Consigliato per Privacy**
- ✅ Completamente locale e privato
- ✅ Molto più accurato
- ✅ Supporta 99+ lingue
- ✅ Gestisce meglio accenti e rumori
- ❌ Richiede installazione: `pip install openai-whisper`
- ❌ Leggermente più lento

**Modelli Whisper:**
- `tiny` (39 MB): Velocissimo, accuratezza base
- `base` (74 MB): Ottimo per CPU
- `small` (244 MB): Buon compromesso
- `medium` (769 MB): **Consigliato per GPU** - Ottimo rapporto qualità/velocità
- `large` (1.5 GB): **Massima accuratezza con RTX 3090** - Veloce su GPU potenti!

### 🎮 Accelerazione GPU
Con la tua RTX 3090:
- Whisper `large` è velocissimo (~2-3 secondi per audio)
- FP16 automatico per prestazioni ottimali
- Raccomandato: usa `medium` o `large`

### 📝 Documento Summary
Alla fine della conversazione, clicca **"📝 Genera Summary (MD + DOCX)"** per:
- Creare automaticamente un documento markdown strutturato
- Riassumere argomenti, decisioni e action items
- Salvare come file .md e .docx scaricabili
- Vedere anteprima prima del download

### 🎤 Input Vocale
1. Clicca sul **microfono** 🎤
2. **Parla** il tuo messaggio
3. Clicca **STOP** ⏹️ per terminare la registrazione
4. ⏳ *L'app trascrive automaticamente...*
5. ✅ **Il testo appare nel campo editabile**
6. 📝 **Puoi modificare/correggere** il testo se necessario
7. Clicca **"📤 Invia"** per mandare a LM Studio
8. 🤖 Ricevi la risposta (scritta e/o vocale)

**Vantaggi:**
- ✅ Vedi cosa ha capito prima di inviare
- ✅ Correggi errori di trascrizione
- ✅ Aggiungi dettagli o modifiche
- ✅ Combina audio + editing manuale
""",
        
        # Status messages
        "no_text": "❌ Nessun testo da inviare. Registra audio e attendi la trascrizione, poi invia.",
        "session_restored": "✅ Ultima sessione caricata",
        "conv_reset": "🔄 Conversazione resettata",
        "response_time": "📊 Tempo risposta",
        "total_messages": "Messaggi totali",
        "avg_time": "Media",
        "playing_audio": "🔊 Riproduzione audio...",
        "response_received": "✅ Risposta ricevuta",
        "no_conversation": "❌ Nessuna conversazione da",
        "json_saved": "✅ Conversazione esportata in",
        "txt_saved": "✅ Conversazione esportata in",
        "summary_error": "❌ Errore nella generazione del riepilogo",
        "connected": "✅ Connesso",
        "not_connected": "❌ Impossibile connettersi a LM Studio",
        "models_available": "modelli disponibili",
        "model_loaded": "✅ Modello caricato",
        "cached": "✅ Caricato dalla cache",
        "downloaded": "✅ Download completato e salvato in cache",
        "cache_location": "📁 Cache",
    }
}

# Detect browser language
def get_browser_language():
    """Detect browser language from request headers"""
    try:
        # This will be set by Gradio from Accept-Language header
        return "it" if "it" in os.environ.get("HTTP_ACCEPT_LANGUAGE", "").lower() else "en"
    except:
        return "en"

# Current language (default English, will be updated from browser)
current_lang = get_browser_language()

def t(key):
    """Translation helper function"""
    return TRANSLATIONS.get(current_lang, TRANSLATIONS["en"]).get(key, key)

# Configurazione globale
LM_STUDIO_URL = "http://localhost:1234/v1/chat/completions"
conversation_history = []
stats = {"messages": 0, "total_time": 0}
stop_speaking = False

# Directory cache per modelli Whisper (evita ri-download)
WHISPER_CACHE_DIR = os.path.join(os.path.expanduser("~"), ".cache", "whisper")
os.makedirs(WHISPER_CACHE_DIR, exist_ok=True)

# File per persistenza conversazioni
SESSIONS_DIR = os.path.join(os.path.dirname(__file__), "sessions")
os.makedirs(SESSIONS_DIR, exist_ok=True)
CURRENT_SESSION_FILE = os.path.join(SESSIONS_DIR, "sessione_corrente.json")

# Carica l'ultima sessione se esiste
def load_last_session():
    """Carica l'ultima sessione salvata"""
    global conversation_history, stats
    if os.path.exists(CURRENT_SESSION_FILE):
        try:
            with open(CURRENT_SESSION_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                conversation_history = data.get('history', [])
                stats = data.get('stats', {"messages": 0, "total_time": 0})
                print(f"✅ Sessione caricata: {len(conversation_history)} messaggi")
                return True
        except Exception as e:
            print(f"⚠️ Errore caricamento sessione: {e}")
    return False

def save_current_session():
    """Salva la sessione corrente"""
    try:
        with open(CURRENT_SESSION_FILE, 'w', encoding='utf-8') as f:
            json.dump({
                'history': conversation_history,
                'stats': stats,
                'timestamp': datetime.now().isoformat()
            }, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"❌ Errore salvataggio sessione: {e}")
        return False

def convert_history_to_chatbot_format(history):
    """Converte conversation_history in formato Gradio chatbot"""
    chatbot_history = []
    i = 0
    while i < len(history):
        msg = history[i]
        if msg.get('role') == 'user':
            user_msg = msg.get('content', '')
            # Cerca il messaggio assistant successivo
            assistant_msg = ''
            if i + 1 < len(history) and history[i + 1].get('role') == 'assistant':
                assistant_msg = history[i + 1].get('content', '')
                i += 2
            else:
                i += 1
            chatbot_history.append({'role': 'user', 'content': user_msg})
            if assistant_msg:
                chatbot_history.append({'role': 'assistant', 'content': assistant_msg})
        else:
            i += 1
    return chatbot_history

# Carica l'ultima sessione all'avvio
load_last_session()

# Inizializza il motore text-to-speech
engine = pyttsx3.init()
voices = engine.getProperty('voices')

# Inizializza il riconoscimento vocale
recognizer = sr.Recognizer()

def load_whisper_model(model_size="base", use_gpu=True):
    """Carica il modello Whisper (usa cache se già scaricato)"""
    global whisper_model
    if not WHISPER_AVAILABLE:
        return "❌ Whisper non installato. Installa con: pip install openai-whisper"
    
    try:
        device = "cuda" if (use_gpu and CUDA_AVAILABLE) else "cpu"
        
        # Whisper scarica automaticamente in ~/.cache/whisper se non presente
        # Non serve fare nulla di speciale, usa già il caching!
        whisper_model = whisper.load_model(model_size, device=device, download_root=WHISPER_CACHE_DIR)
        
        gpu_info = f" su {GPU_NAME}" if device == "cuda" else " su CPU"
        cache_info = "✅ Modello già in cache" if os.path.exists(os.path.join(WHISPER_CACHE_DIR, f"{model_size}.pt")) else "⬇️ Download completato"
        return f"{cache_info}\n✅ Modello Whisper '{model_size}' caricato{gpu_info}!"
    except Exception as e:
        return f"❌ Errore nel caricamento: {str(e)}"

def get_available_voices():
    """Ottiene le voci disponibili sul sistema"""
    return {f"{v.name} ({v.languages[0] if v.languages else 'unknown'})": i 
            for i, v in enumerate(voices)}

def set_voice(voice_index, rate, volume):
    """Configura la voce del TTS"""
    engine.setProperty('voice', voices[voice_index].id)
    engine.setProperty('rate', rate)
    engine.setProperty('volume', volume)

def parla(testo, rate, volume, voice_index):
    """Converte il testo in voce"""
    global stop_speaking
    stop_speaking = False
    set_voice(voice_index, rate, volume)
    
    # Dividi in frasi per permettere l'interruzione
    frasi = testo.replace('!', '.').replace('?', '.').split('.')
    for frase in frasi:
        if stop_speaking:
            break
        if frase.strip():
            engine.say(frase.strip())
            engine.runAndWait()

def stop_speech():
    """Interrompe la sintesi vocale"""
    global stop_speaking
    stop_speaking = True
    engine.stop()

def ascolta_audio(audio_data, recognition_method="google", whisper_lang="it"):
    """Trascrivi audio da Gradio"""
    if audio_data is None:
        return "❌ Nessun audio registrato"
    
    try:
        # Gradio Audio restituisce una tupla (sample_rate, numpy_array) o un file path
        if isinstance(audio_data, tuple):
            sample_rate, audio_array = audio_data
        elif isinstance(audio_data, str):
            # Se è un percorso file, caricalo
            import soundfile as sf
            audio_array, sample_rate = sf.read(audio_data)
        else:
            return "❌ Formato audio non riconosciuto"
        
        # Metodo Whisper (locale)
        if recognition_method == "whisper":
            if not WHISPER_AVAILABLE or whisper_model is None:
                return "❌ Whisper non disponibile. Usa Google o carica il modello."
            
            # Assicurati che l'audio sia mono
            if len(audio_array.shape) > 1:
                audio_array = audio_array.mean(axis=1)
            
            # Converti a float32 e normalizza se necessario
            if audio_array.dtype == np.int16:
                audio_float = audio_array.astype(np.float32) / 32768.0
            elif audio_array.dtype == np.int32:
                audio_float = audio_array.astype(np.float32) / 2147483648.0
            else:
                audio_float = audio_array.astype(np.float32)
            
            # Whisper si aspetta 16kHz, ricampiona se necessario
            if sample_rate != 16000:
                if LIBROSA_AVAILABLE:
                    import librosa
                    audio_float = librosa.resample(audio_float, orig_sr=sample_rate, target_sr=16000)
                else:
                    # Resampling semplice senza librosa
                    from scipy import signal
                    num_samples = int(len(audio_float) * 16000 / sample_rate)
                    audio_float = signal.resample(audio_float, num_samples)
            
            # Trascrivi con Whisper (usa GPU se disponibile)
            result = whisper_model.transcribe(
                audio_float,
                language=whisper_lang,
                fp16=CUDA_AVAILABLE  # Usa FP16 su GPU per velocità
            )
            return result["text"].strip()
        
        # Metodo Google (online - default)
        else:
            # Converti l'array numpy in bytes per SpeechRecognition
            if audio_array.dtype != np.int16:
                # Converti in int16 se necessario
                audio_array = (audio_array * 32768.0).astype(np.int16)
            
            # Assicurati che sia mono
            if len(audio_array.shape) > 1:
                audio_array = audio_array.mean(axis=1).astype(np.int16)
            
            # Crea l'oggetto AudioData
            audio_bytes = audio_array.tobytes()
            sr_audio = sr.AudioData(audio_bytes, sample_rate, 2)
            
            # Mappa lingua per Google
            google_lang_map = {
                "it": "it-IT",
                "en": "en-US",
                "es": "es-ES",
                "fr": "fr-FR",
                "de": "de-DE"
            }
            google_lang = google_lang_map.get(whisper_lang, "it-IT")
            
            testo = recognizer.recognize_google(sr_audio, language=google_lang)
            return testo
            
    except sr.UnknownValueError:
        return "❌ Non ho capito, riprova"
    except Exception as e:
        return f"❌ Errore: {str(e)}"

def ascolta_microfono():
    """Ascolta dal microfono in tempo reale"""
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source, duration=0.5)
        try:
            audio = recognizer.listen(source, timeout=5, phrase_time_limit=15)
            testo = recognizer.recognize_google(audio, language="it-IT")
            return testo
        except sr.WaitTimeoutError:
            return "❌ Timeout - nessun audio rilevato"
        except sr.UnknownValueError:
            return "❌ Non ho capito, riprova"
        except Exception as e:
            return f"❌ Errore: {str(e)}"

def get_available_models():
    """Ottiene i modelli disponibili da LM Studio"""
    try:
        response = requests.get("http://localhost:1234/v1/models", timeout=5)
        if response.status_code == 200:
            models = response.json().get('data', [])
            model_list = [m['id'] for m in models] if models else []
            # Aggiungi sempre "Auto" come prima opzione
            if model_list:
                return ["Auto (usa modello caricato)"] + model_list
            return ["Auto (usa modello caricato)"]
        return ["Auto (usa modello caricato)"]
    except:
        return ["Auto (usa modello caricato)"]

def chiedi_a_lm_studio(messaggio, system_prompt, temperature, max_tokens, model, lingua):
    """Invia una richiesta a LM Studio"""
    global conversation_history, stats
    
    start_time = datetime.now()
    
    # Aggiungi il messaggio allo storico
    conversation_history.append({"role": "user", "content": messaggio})
    
    # Prepara i messaggi con system prompt - FORMATO CORRETTO
    messages = [{"role": "system", "content": system_prompt}]
    
    # Aggiungi solo messaggi con formato corretto
    for msg in conversation_history:
        if isinstance(msg, dict) and "role" in msg and "content" in msg:
            messages.append({"role": msg["role"], "content": msg["content"]})
    
    try:
        payload = {
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        
        # Aggiungi model solo se non è "Auto"
        if model and not model.startswith("Auto"):
            payload["model"] = model
        
        risposta = requests.post(
            LM_STUDIO_URL,
            json=payload,
            timeout=60
        )
        
        if risposta.status_code == 200:
            risultato = risposta.json()
            testo_risposta = risultato['choices'][0]['message']['content']
            conversation_history.append({"role": "assistant", "content": testo_risposta})
            
            # Aggiorna statistiche
            elapsed = (datetime.now() - start_time).total_seconds()
            stats["messages"] += 1
            stats["total_time"] += elapsed
            
            return testo_risposta, elapsed
        else:
            return f"❌ Errore {risposta.status_code}: {risposta.text}", 0
            
    except requests.exceptions.ConnectionError:
        return "❌ Impossibile connettersi a LM Studio. Assicurati che sia in esecuzione sulla porta 1234.", 0
    except Exception as e:
        return f"❌ Errore: {str(e)}", 0

def update_connection_status():
    """Aggiorna l'indicatore di connessione"""
    try:
        response = requests.get("http://localhost:1234/v1/models", timeout=2)
        if response.status_code == 200:
            models = response.json().get('data', [])
            return f"""
                <div style="padding: 10px; border-radius: 8px; background: #065f46; text-align: center;">
                    <span style="color: #10b981; font-size: 20px;">●</span>
                    <span style="color: #d1fae5; margin-left: 8px;">Connesso - {len(models)} modelli</span>
                </div>
            """
        return """
            <div style="padding: 10px; border-radius: 8px; background: #92400e; text-align: center;">
                <span style="color: #fbbf24; font-size: 20px;">●</span>
                <span style="color: #fef3c7; margin-left: 8px;">Connesso - Nessun modello</span>
            </div>
        """
    except:
        return """
            <div style="padding: 10px; border-radius: 8px; background: #7f1d1d; text-align: center;">
                <span style="color: #ef4444; font-size: 20px;">●</span>
                <span style="color: #fee2e2; margin-left: 8px;">Non connesso</span>
            </div>
        """

def trascrivi_audio_automatico(audio, recognition_method, whisper_lang):
    """Trascrivi automaticamente quando l'audio è registrato"""
    if audio is None:
        return ""
    
    # Estrai solo il metodo (rimuovi la descrizione)
    method = recognition_method.split(" - ")[0]
    
    # Trascrivi
    testo = ascolta_audio(audio, method, whisper_lang)
    
    if testo.startswith("❌"):
        return testo
    
    return testo

def processa_messaggio(audio, testo_input, system_prompt, temperature, max_tokens, 
                       model, speak_output, rate, volume, voice_index, lingua, history,
                       recognition_method, whisper_lang):
    """Processa il messaggio vocale o testuale"""
    # Aggiorna status connessione
    conn_status = update_connection_status()
    
    # Assicurati che history sia una lista (non None)
    if history is None:
        history = []
    
    # Usa il testo se disponibile, altrimenti non fare nulla
    if not testo_input or not testo_input.strip():
        return history, "", "❌ Nessun testo da inviare. Registra audio e attendi la trascrizione, poi invia.", "", conn_status
    
    messaggio = testo_input.strip()
    
    # FORMATO DICT per Gradio 6 (messages format)
    new_history = list(history) if history else []
    
    # Ottieni risposta
    risposta, elapsed = chiedi_a_lm_studio(messaggio, system_prompt, temperature, 
                                           max_tokens, model, lingua)
    
    # Aggiungi messaggi come dict separati
    new_history.append({"role": "user", "content": messaggio})
    new_history.append({"role": "assistant", "content": risposta})
    
    # Salva sessione automaticamente
    save_current_session()
    
    # Info statistiche
    if stats['messages'] > 0:
        stats_text = f"📊 Tempo risposta: {elapsed:.2f}s | Messaggi totali: {stats['messages']} | Tempo medio: {stats['total_time']/stats['messages']:.2f}s"
    else:
        stats_text = f"📊 Tempo risposta: {elapsed:.2f}s | Messaggi totali: 1"
    
    # Pronuncia la risposta se richiesto
    if speak_output and not risposta.startswith("❌"):
        threading.Thread(target=parla, args=(risposta, rate, volume, voice_index), daemon=True).start()
        status = "🔊 Riproduzione audio..."
    else:
        status = "✅ Risposta ricevuta"
    
    return new_history, "", status, stats_text, conn_status

def clear_conversation():
    """Pulisce la conversazione"""
    global conversation_history, stats
    conversation_history = []
    stats = {"messages": 0, "total_time": 0}
    save_current_session()  # Salva lo stato vuoto
    conn_status = update_connection_status()
    # Ritorna lista vuota (non tuple)
    return [], "", "🔄 Conversazione resettata", "", conn_status

def save_conversation():
    """Salva la conversazione in un file"""
    if not conversation_history:
        return "❌ Nessuna conversazione da salvare"
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"conversation_{timestamp}.json"
    
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(conversation_history, f, ensure_ascii=False, indent=2)
    
    return f"✅ Conversazione salvata in: {filename}"

def export_as_text():
    """Esporta la conversazione come testo"""
    if not conversation_history:
        return "❌ Nessuna conversazione da esportare"
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"conversation_{timestamp}.txt"
    
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(f"=== Conversazione del {datetime.now().strftime('%d/%m/%Y %H:%M')} ===\n\n")
        for msg in conversation_history:
            role = "TU" if msg["role"] == "user" else "ASSISTENTE"
            f.write(f"{role}: {msg['content']}\n\n")
    
    return f"✅ Conversazione esportata in: {filename}"

def export_document_as_md(content, title="Documento"):
    """Esporta contenuto come Markdown"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{title.lower().replace(' ', '_')}_{timestamp}.md"
    
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            f.write(f"**Data**: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n")
            f.write("---\n\n")
            f.write(content)
        return filename, f"✅ Documento MD salvato: {filename}"
    except Exception as e:
        return None, f"❌ Errore: {str(e)}"

def export_document_as_docx(content, title="Documento"):
    """Esporta contenuto come DOCX"""
    if not DOCX_AVAILABLE:
        return None, "❌ python-docx non installato. Installa con: pip install python-docx"
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{title.lower().replace(' ', '_')}_{timestamp}.docx"
    
    try:
        doc = Document()
        
        # Titolo
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Data
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_run.italic = True
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Separatore
        doc.add_paragraph("_" * 50)
        
        # Contenuto - parse markdown semplice
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Headers
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # Liste
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and '. ' in line[:4]:
                doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            # Testo normale
            else:
                doc.add_paragraph(line)
        
        doc.save(filename)
        return filename, f"✅ Documento DOCX salvato: {filename}"
    except Exception as e:
        return None, f"❌ Errore: {str(e)}"

def generate_summary_document():
    """Genera un documento markdown riassuntivo della conversazione"""
    if not conversation_history:
        return None, None, "❌ Nessuna conversazione da riassumere"
    
    # Prepara il prompt per il riassunto
    conv_text = "\n".join([f"{msg['role']}: {msg['content']}" for msg in conversation_history])
    
    summary_prompt = f"""Basandoti su questa conversazione, crea un documento markdown strutturato e professionale che riassuma:

1. **Argomento principale** discusso
2. **Punti chiave** emersi
3. **Decisioni** o conclusioni raggiunte
4. **Action items** o prossimi passi (se presenti)
5. **Dettagli tecnici** importanti menzionati

Conversazione:
{conv_text}

Crea un documento markdown ben formattato con titoli, elenchi puntati e sezioni chiare. NON includere preamble, rispondi SOLO con il markdown del documento."""

    try:
        response = requests.post(
            LM_STUDIO_URL,
            json={
                "messages": [
                    {"role": "system", "content": "Sei un assistente esperto nella creazione di documenti riassuntivi chiari e ben strutturati in formato markdown. Rispondi SOLO con il markdown, senza preamble."},
                    {"role": "user", "content": summary_prompt}
                ],
                "temperature": 0.3,
                "max_tokens": 2000
            },
            timeout=60
        )
        
        if response.status_code == 200:
            risultato = response.json()
            summary = risultato['choices'][0]['message']['content']
            
            # Salva MD
            md_file, md_status = export_document_as_md(summary, "Summary")
            
            # Salva DOCX
            docx_file, docx_status = export_document_as_docx(summary, "Summary")
            
            status_msg = f"{md_status}\n{docx_status if docx_file else '⚠️ DOCX non disponibile (installa python-docx)'}"
            
            return summary, md_file, status_msg
        else:
            return None, None, f"❌ Errore nella generazione: {response.status_code}"
            
    except Exception as e:
        return None, None, f"❌ Errore: {str(e)}"

# Crea l'interfaccia Gradio con tema semplice
with gr.Blocks(title="Chat Vocale LM Studio") as demo:
    gr.Markdown("""
    # 🎙️ Chat Vocale con LM Studio
    ### Interagisci con i modelli AI tramite voce o testo
    """)
    
    with gr.Tabs():
        # Tab principale - Chat
        with gr.Tab("💬 Chat"):
            with gr.Row():
                with gr.Column(scale=2):
                    chatbot = gr.Chatbot(
                        label="Conversazione",
                        height=500
                    )
                    
                    with gr.Row():
                        with gr.Column(scale=1):
                            audio_input = gr.Audio(
                                sources=["microphone"],
                                type="numpy",
                                label="🎤 Registra Audio",
                                format="wav"
                            )
                        with gr.Column(scale=2):
                            testo_input = gr.Textbox(
                                label="✍️ Testo (auto-compilato da audio o scrivi manualmente)",
                                placeholder="Registra audio o scrivi qui...",
                                lines=3
                            )
                    
                    with gr.Row():
                        invia_btn = gr.Button("📤 Invia", variant="primary", size="lg")
                        stop_btn = gr.Button("⏹️ Ferma Audio", size="lg")
                        clear_btn = gr.Button("🗑️ Pulisci Chat", size="lg")
                    
                    status_box = gr.Textbox(label="📊 Status", interactive=False)
                    stats_box = gr.Textbox(label="📈 Statistiche", interactive=False)
                
                with gr.Column(scale=1):
                    gr.Markdown("### ⚙️ Impostazioni Rapide")
                    
                    # Indicatore connessione LM Studio
                    connection_indicator = gr.HTML("""
                        <div style="padding: 10px; border-radius: 8px; background: #1e293b; text-align: center;">
                            <span style="color: #10b981; font-size: 20px;">●</span>
                            <span style="color: #e2e8f0; margin-left: 8px;">Verifica connessione...</span>
                        </div>
                    """)
                    
                    # Impostazioni riconoscimento vocale
                    with gr.Accordion("🎤 Riconoscimento Vocale", open=True):
                        recognition_method = gr.Radio(
                            choices=[
                                "google - Online (veloce, richiede internet)",
                                "whisper - Locale (più accurato, privacy)"
                            ],
                            value="google - Online (veloce, richiede internet)",
                            label="Metodo di Riconoscimento",
                            info="Whisper è più accurato ma richiede installazione"
                        )
                        
                        whisper_lang = gr.Dropdown(
                            choices=["it", "en", "es", "fr", "de", "pt", "ru", "ja", "zh"],
                            value="it",
                            label="🌍 Lingua Audio Input"
                        )
                        
                        if WHISPER_AVAILABLE:
                            whisper_status = gr.Markdown("✅ **Whisper disponibile**")
                            if CUDA_AVAILABLE:
                                gr.Markdown(f"🎮 **GPU Accelerazione Attiva**: {GPU_NAME}")
                            with gr.Row():
                                whisper_model_size = gr.Dropdown(
                                    choices=["tiny", "base", "small", "medium", "large"],
                                    value="medium" if CUDA_AVAILABLE else "base",
                                    label="Modello Whisper",
                                    info="tiny=39MB, base=74MB, small=244MB, medium=769MB, large=1.5GB (consigliato con GPU!)"
                                )
                                use_gpu_whisper = gr.Checkbox(
                                    label="🎮 Usa GPU",
                                    value=CUDA_AVAILABLE,
                                    interactive=CUDA_AVAILABLE
                                )
                                load_whisper_btn = gr.Button("📥 Carica Modello", size="sm")
                            whisper_load_status = gr.Textbox(label="Status", interactive=False)
                        else:
                            whisper_status = gr.Markdown(
                                "⚠️ **Whisper non installato**\n\n"
                                "Per usare il riconoscimento locale:\n"
                                "```bash\n"
                                "pip install openai-whisper\n"
                                "```"
                            )
                    
                    speak_output = gr.Checkbox(label="🔊 Leggi risposta ad alta voce", value=True)
                    
                    lingua = gr.Dropdown(
                        choices=["Italiano", "English", "Español", "Français"],
                        value="Italiano",
                        label="🌍 Lingua"
                    )
                    
                    with gr.Accordion("🎛️ Parametri Modello", open=False):
                        model_dropdown = gr.Dropdown(
                            choices=get_available_models(),
                            value="Auto (usa modello caricato)",
                            label="🤖 Modello",
                            info="Seleziona il modello da LM Studio"
                        )
                        
                        temperature = gr.Slider(0, 2, value=0.7, step=0.1, 
                                               label="🌡️ Temperature (Creatività)")
                        max_tokens = gr.Slider(50, 2000, value=500, step=50,
                                              label="📏 Max Tokens (Lunghezza)")
                    
                    with gr.Accordion("🎨 Personalizzazione Risposta", open=False):
                        system_prompt = gr.Textbox(
                            label="📝 System Prompt",
                            value="Sei un assistente utile che risponde in italiano in modo conciso e chiaro.",
                            lines=3
                        )
                    
                    with gr.Accordion("🔊 Impostazioni Voce", open=False):
                        voice_dropdown = gr.Dropdown(
                            choices=list(get_available_voices().keys()),
                            value=list(get_available_voices().keys())[0] if voices else None,
                            label="🎭 Voce"
                        )
                        rate = gr.Slider(50, 300, value=150, step=10, label="⚡ Velocità")
                        volume = gr.Slider(0, 1, value=0.9, step=0.1, label="🔊 Volume")
                    
                    with gr.Accordion("💾 Esporta Conversazione", open=False):
                        save_json_btn = gr.Button("💾 Salva JSON")
                        save_txt_btn = gr.Button("📄 Esporta TXT")
                        generate_summary_btn = gr.Button("📝 Genera Summary (MD + DOCX)", variant="primary")
                        save_status = gr.Textbox(label="Status Salvataggio", interactive=False)
                        
                    with gr.Accordion("📄 Anteprima Summary", open=False, visible=False) as summary_accordion:
                        summary_preview = gr.Markdown(label="Summary Generated")
                        download_info = gr.Markdown("ℹ️ **File salvati** - Controlla la cartella del progetto per i file .md e .docx")
        
        # Tab Impostazioni Avanzate
        with gr.Tab("🔧 Impostazioni Avanzate"):
            gr.Markdown("### 🌐 Configurazione LM Studio")
            lm_url = gr.Textbox(value=LM_STUDIO_URL, label="URL Server LM Studio")
            test_connection = gr.Button("🔌 Test Connessione")
            connection_status = gr.Textbox(label="Status Connessione", interactive=False)
            
            gr.Markdown("### 🎮 Info GPU")
            if CUDA_AVAILABLE:
                gr.Markdown(f"""
                ✅ **GPU Rilevata**: {GPU_NAME}
                - **CUDA disponibile**: Sì
                - **Accelerazione Whisper**: Attiva
                - **Performance**: ~10-20x più veloce su GPU!
                
                **Raccomandazioni per RTX 3090:**
                - Usa modello Whisper `large` per massima accuratezza
                - Abilita GPU acceleration in LM Studio
                - FP16 automatico attivo per velocità
                """)
            else:
                gr.Markdown("⚠️ **GPU non rilevata** - Utilizzo CPU")
            
            gr.Markdown("### 📋 Info Sistema")
            gr.Markdown(f"- **Voci disponibili:** {len(voices)}")
            gr.Markdown(f"- **Riconoscimento vocale:** Google Speech Recognition + Whisper")
            gr.Markdown(f"- **Sintesi vocale:** pyttsx3")
            gr.Markdown(f"- **Whisper disponibile:** {'✅ Sì' if WHISPER_AVAILABLE else '❌ No'}")
        
        # Tab Info
        with gr.Tab("ℹ️ Guida"):
            gr.Markdown("""
            ## 🎯 Come usare l'applicazione
            
            ### 🎤 Riconoscimento Vocale
            
            **Google (Online) - Default**
            - ✅ Veloce e affidabile
            - ✅ Non richiede installazione
            - ❌ Richiede connessione internet
            - ❌ I dati audio vengono inviati a Google
            
            **Whisper (Locale) - Consigliato per Privacy**
            - ✅ Completamente locale e privato
            - ✅ Molto più accurato
            - ✅ Supporta 99+ lingue
            - ✅ Gestisce meglio accenti e rumori
            - ❌ Richiede installazione: `pip install openai-whisper`
            - ❌ Leggermente più lento
            
            **Modelli Whisper:**
            - `tiny` (39 MB): Velocissimo, accuratezza base
            - `base` (74 MB): Ottimo per CPU
            - `small` (244 MB): Buon compromesso
            - `medium` (769 MB): **Consigliato per GPU** - Ottimo rapporto qualità/velocità
            - `large` (1.5 GB): **Massima accuratezza con RTX 3090** - Veloce su GPU potenti!
            
            ### 🎮 Accelerazione GPU
            Con la tua RTX 3090:
            - Whisper `large` è velocissimo (~2-3 secondi per audio)
            - FP16 automatico per prestazioni ottimali
            - Raccomandato: usa `medium` o `large`
            
            ### 📝 Documento Riepilogo
            Alla fine della conversazione, clicca su **"📝 Genera Documento Riepilogo"** per:
            - Creare automaticamente un documento markdown strutturato
            - Riassumere argomenti, decisioni e action items
            - Salvare come file .md scaricabile
            - Vedere anteprima prima del download
            
            ### 🎤 Input Vocale
            1. Clicca sul pulsante microfono
            2. Parla chiaramente
            3. Clicca su "Invia"
            
            ### ✍️ Input Testuale
            1. Scrivi nella casella di testo
            2. Clicca su "Invia"
            
            ### ⚙️ Impostazioni
            - **Temperature**: Controlla la creatività (0 = prevedibile, 2 = creativo)
            - **Max Tokens**: Lunghezza massima della risposta
            - **System Prompt**: Definisce il comportamento dell'assistente
            - **Voce/Velocità/Volume**: Personalizza l'output audio
            
            ### 💡 Suggerimenti
            - Parla in un ambiente silenzioso per miglior riconoscimento
            - Usa "Ferma Audio" per interrompere risposte lunghe
            - Salva le conversazioni interessanti per riferimento futuro
            - Regola la temperature in base al tipo di risposta desiderata
            
            ### ⚠️ Risoluzione Problemi
            - **"Impossibile connettersi"**: Verifica che LM Studio sia avviato
            - **"Non ho capito"**: Parla più chiaramente o usa l'input testuale
            - **Audio non funziona**: Controlla le impostazioni del microfono
            """)
    
    # Collegamenti eventi
    voice_map = get_available_voices()
    
    # Trascrizione automatica quando finisce la registrazione audio
    audio_input.stop_recording(
        fn=trascrivi_audio_automatico,
        inputs=[audio_input, recognition_method, whisper_lang],
        outputs=[testo_input]
    )
    
    def process_with_recognition(a, t, sp, temp, mt, m, so, r, v, voice, l, h, rec_method, w_lang):
        # Estrai solo il nome del metodo (prima del trattino)
        method = rec_method.split(" - ")[0]
        return processa_messaggio(a, t, sp, temp, mt, m, so, r, v, voice_map[voice], l, h, method, w_lang)
    
    invia_btn.click(
        fn=process_with_recognition,
        inputs=[audio_input, testo_input, system_prompt, temperature, max_tokens,
                model_dropdown, speak_output, rate, volume, voice_dropdown, lingua, chatbot,
                recognition_method, whisper_lang],
        outputs=[chatbot, testo_input, status_box, stats_box, connection_indicator]
    )
    
    stop_btn.click(fn=stop_speech, outputs=[])
    clear_btn.click(fn=clear_conversation, outputs=[chatbot, testo_input, status_box, stats_box, connection_indicator])
    
    save_json_btn.click(fn=save_conversation, outputs=[save_status])
    save_txt_btn.click(fn=export_as_text, outputs=[save_status])
    
    def handle_summary_generation():
        summary, md_file, status = generate_summary_document()
        if summary:
            return summary, status, gr.Accordion(visible=True, open=True)
        return "", status, gr.Accordion(visible=False)
    
    generate_summary_btn.click(
        fn=handle_summary_generation,
        outputs=[summary_preview, save_status, summary_accordion]
    )
    
    # Caricamento modello Whisper
    if WHISPER_AVAILABLE:
        load_whisper_btn.click(
            fn=load_whisper_model,
            inputs=[whisper_model_size, use_gpu_whisper],
            outputs=[whisper_load_status]
        )
    
    def test_lm_connection():
        try:
            response = requests.get("http://localhost:1234/v1/models", timeout=5)
            if response.status_code == 200:
                models = response.json().get('data', [])
                return f"✅ Connesso! Modelli disponibili: {len(models)}"
            return f"⚠️ Connesso ma nessun modello caricato"
        except:
            return "❌ Impossibile connettersi a LM Studio"
    
    test_connection.click(fn=test_lm_connection, outputs=[connection_status])
    
    # Auto-refresh connessione all'avvio
    demo.load(fn=update_connection_status, outputs=[connection_indicator])
    
    # Carica l'ultima sessione all'avvio dell'interfaccia
    def load_session_on_start():
        if conversation_history:
            chatbot_hist = convert_history_to_chatbot_format(conversation_history)
            stats_text = f"📊 Sessione ripristinata: {len(conversation_history)//2} messaggi | Tempo totale: {stats['total_time']:.1f}s"
            return chatbot_hist, "✅ Ultima sessione caricata", stats_text
        return [], "", ""
    
    demo.load(fn=load_session_on_start, outputs=[chatbot, status_box, stats_box])

if __name__ == "__main__":
    print("🚀 Avvio Chat Vocale con LM Studio...")
    print("📱 L'interfaccia si aprirà nel browser...")
    demo.launch(share=False, inbrowser=True)
    